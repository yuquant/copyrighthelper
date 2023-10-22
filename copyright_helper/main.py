import os
import codecs
import tempfile
from .utils import remove_annotations, Extractor

__all__ = ['write_cleaned_code_to_doc', 'CopyrightHelper', 'remove_annotations']


class CopyrightHelper:
    def __init__(self, code_files: list[str]):
        self._code_files = code_files

    @classmethod
    def from_dir(cls, dir_path):
        file_paths = []
        for root, dirs, files in os.walk(dir_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_paths.append(file_path)
        return cls(code_files=file_paths)

    def get_removed_code(self) -> str:
        codes = []
        lines = 0
        for file in self._code_files:
            if lines <= 3000 and self._is_code_file(file):
                with open(file, 'r') as f:
                    code = f.read()
                code = remove_annotations(code)
                codes.append(code)
                lines += len(code.split('\n'))
        content = '\n'.join(codes)
        res = '\n'.join(content.split('\n')[:3000])

        return res

    def write_to_docx(self, code: str, file_path):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
        # lines = code.split('\n')
        doc = Document()
        p = doc.add_paragraph('')  # 增加一页
        doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
        doc.styles['Normal'].font.size = Pt(8)  # 设置字体的大小为 5 号字体
        p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
        paragraph_format = doc.styles['Normal'].paragraph_format
        paragraph_format.line_spacing = Pt(12.9)  # 固定值12,9磅, 保证每页有50行代码
        # for line in lines[:3000]:
        #     p.add_run(line)
        p.add_run(code)
        doc.save(file_path)  # 不足60 页进行保存

    def _is_code_file(self, file_path: str) -> bool:
        file_name = os.path.split(file_path)[1]
        # suffix = file_name.split('.')
        if file_name.startswith('.'):
            return False
        # elif suffix[-1] not in ['txt', 'py', 'js', 'cpp', 'cs', 'html', 'java', 'rb', 'php', 'kt', 'm', 'c', 'css']:
        #     return False
        _TEXT_BOMS = (
            codecs.BOM_UTF16_BE,
            codecs.BOM_UTF16_LE,
            codecs.BOM_UTF32_BE,
            codecs.BOM_UTF32_LE,
            codecs.BOM_UTF8,
        )

        with open(file_path, 'rb') as file:
            initial_bytes = file.read(8192)
            file.close()
        is_binary = not any(initial_bytes.startswith(bom) for bom in _TEXT_BOMS) and b'\0' in initial_bytes
        return not is_binary


def write_cleaned_code_to_doc(code_src: str, doc_path):
    """

    :param code_src: txt,zip,tar,directory is supported
    :param doc_path: output docx file path
    :return:
    """
    dp = tempfile.TemporaryDirectory()
    if os.path.isdir(code_src):
        helper = CopyrightHelper.from_dir(code_src)
    elif os.path.isfile(code_src):
        extractor = Extractor(code_src)
        extractor.extract_to(dp.name)
        helper = CopyrightHelper.from_dir(dp.name)
    else:
        raise ValueError('Not a file or directory')
    code = helper.get_removed_code()
    dp.cleanup()
    helper.write_to_docx(code, doc_path)

